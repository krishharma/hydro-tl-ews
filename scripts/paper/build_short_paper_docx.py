#!/usr/bin/env python3
"""Build the short paper as a Word manuscript from poster_metrics.json.

Primary compiled artifact (replaces the LaTeX PDF as the shareable paper)::

    docs/Short_Paper_Hydro_TL_EWS.docx

Run after ``scripts/analysis/poster_from_csvs.py``.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
METRICS = ROOT / "results" / "midwest_mini" / "poster" / "poster_metrics.json"
FIG = ROOT / "docs" / "paper_images"
OUT = ROOT / "docs" / "Short_Paper_Hydro_TL_EWS.docx"

TEAL = RGBColor(0x01, 0x69, 0x6F)
INK = RGBColor(0x1F, 0x1E, 0x1B)
MUTED = RGBColor(0x7A, 0x79, 0x74)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_run(run, *, size=11, bold=False, italic=False, color=INK, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_border(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "D4D0C8")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_para(doc, text, *, size=11, bold=False, italic=False, space_after=8,
             space_before=0, align="left", color=INK, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    _set_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run(run, size=14 if level == 1 else 12, bold=True, color=TEAL)
    return p


def add_caption(doc, text):
    p = add_para(doc, text, size=9, italic=True, space_after=12, space_before=4,
                 color=MUTED, first_line=False)
    return p


def add_figure(doc, path: Path, caption: str, width=6.3):
    if not path.exists():
        add_para(doc, f"[Missing figure: {path.name}]", italic=True, color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True
    hdr = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = ""
        p = hdr[j].paragraphs[0]
        run = p.add_run(h)
        _set_run(run, size=9, bold=True, color=WHITE)
        _shade_cell(hdr[j], "01696F")
        _set_cell_border(hdr[j])
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            _set_run(run, size=9, bold=(j == 0), color=INK)
            if i % 2 == 1:
                _shade_cell(cell, "F4F1EA")
            _set_cell_border(cell)
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def _fmt(x, nd=3):
    if x is None:
        return "—"
    try:
        x = float(x)
    except (TypeError, ValueError):
        return "—"
    if x != x:
        return "—"
    return f"{x:.{nd}f}"


def _pm(x, nd=1):
    v = float(x)
    sign = "+" if v > 0 else ""
    return f"{sign}{v:.{nd}f}"


def build() -> Path:
    m = json.loads(METRICS.read_text())
    skill = m["skill"]
    ews = m["early_warning"]
    prune = m["pruning"]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    add_para(doc,
             "How few donor basins are enough?",
             size=22, bold=True, align="center", space_after=4, color=TEAL)
    add_para(doc,
             "Regional transfer learning for daily streamflow with two years of local data",
             size=13, italic=True, align="center", space_after=8, color=TEAL)
    add_para(doc, "Krish Sharma  ·  August 2026",
             size=11, align="center", space_after=2, color=MUTED)
    add_para(doc, "Code: github.com/krishharma/hydro-tl-ews",
             size=10, align="center", space_after=14, color=MUTED)

    add_heading_custom(doc, "Abstract")
    add_para(
        doc,
        "How many similar donor basins does an entity-aware LSTM need when the "
        "local record is only about two years long and the hardware is a laptop? "
        "Does regional transfer still beat training from scratch? The completed "
        "study is a Midwest CAMELS case at Lick Creek near Perry, Missouri "
        "(USGS 05507600). Seventeen hydrologically similar donors, a 90-day "
        "lookback, and walk-forward evaluation on 2011–2014 are the design, not "
        "a substitute for a continental run. Transfer learning outperformed a "
        "local-from-scratch LSTM and two naive baselines. Walk-forward "
        "Nash–Sutcliffe efficiency (NSE) is 0.179, versus −0.166 for next-day "
        "persistence, −2.34 for a day-of-year climatology fit on the same "
        "two-year warmup, and 0.009 for the local LSTM. Cutting the donor pool "
        "from 17 to 8 basins made skill worse. Flood Q95 warnings rank well "
        "(AUC ≈ 0.91–0.92), but independence-based multi-day probabilities are "
        "overconfident; replacing that product with a window-max of 1-day "
        "probabilities brings Brier scores back near 0.10–0.12. Low-flow "
        "(Q5 = 0 mm d−1) detection is reported separately from drought.",
        align="justify", space_after=10,
    )

    add_heading_custom(doc, "1.  Introduction")
    add_para(
        doc,
        "Most gauges do not come with decades of discharge. The usual large-sample "
        "LSTM recipe—hundreds of basins, long sequences, and a GPU—answers a "
        "different question from the one a field office actually faces: can a "
        "regional model help when the local record is about two years long and "
        "the hardware is ordinary? That constraint is the study design here.",
        align="justify",
    )
    add_para(
        doc,
        "The experiment holds the target, the warmup, and the evaluation window "
        "fixed, and varies only how the regional prior is built. An entity-aware "
        "LSTM (EA-LSTM) is pretrained on similar Midwest donors, then transferred "
        "to Lick Creek with a frozen-LSTM, head-only fine-tune and a walk-forward "
        "refit protocol. Training from scratch on the same warmup is the local "
        "baseline. Next-day persistence and a day-of-year climatology—fit only on "
        "that warmup—put a floor under the NSE numbers. A later sensitivity run "
        "cuts the donor pool to eight similar basins. The numbers in "
        "results/midwest_mini/ are from this completed Lick Creek study "
        "(1,461 evaluation days), not from a smoke test or a synthetic sample.",
        align="justify",
    )

    add_heading_custom(doc, "2.  Data and methods")
    add_heading_custom(doc, "2.1  Basin, donors, and periods", level=2)
    add_para(
        doc,
        "CAMELS-US supplies Daymet weather, USGS streamflow, and 27 static "
        "catchment attributes (Addor et al., 2017; Kratzert et al., 2019). The "
        "target is Lick Creek near Perry, MO (05507600), a plains basin with a "
        "snow fraction of about 0.07. Donors are the Midwest gauges most similar "
        "to the target in static attribute space, restricted to "
        "[36.5°N, 49.5°N] × [−104°W, −80.5°W], after a 50 km exclusion buffer "
        "and after dropping gauges that are missing from the local "
        "HUC 04/05/07/10 extract. That procedure yields 17 donors (Figure 1). "
        "Weather inputs are precipitation, Tmax, Tmin, shortwave radiation, "
        "vapor pressure, and day length. Streamflow is converted to mm d−1.",
        align="justify",
    )
    add_figure(
        doc, FIG / "fig_donor_target_map.png",
        "Figure 1. Seventeen similar Midwest donor gauges and the Lick Creek target. "
        "Grey points are other CAMELS gauges in the same geographic box; they are "
        "context, not training data.",
    )
    add_table(
        doc,
        ["Period", "Dates", "Role"],
        [
            ["Donor pretrain", "Oct 1995 – Sep 2008", "Train the regional EA-LSTM"],
            ["Pretrain validation", "Oct 2008 – Sep 2010", "Early stopping"],
            ["Target warmup", "Jan 2009 – Dec 2010", "Scarce local record (~2 years)"],
            ["Target evaluation", "2011 – 2014", "Fixed-window and walk-forward tests"],
            ["Threshold history", "1990 – 2010", "At-site Q5, Q95, Q99"],
        ],
    )
    add_caption(doc, "Table 1. Time periods used for USGS 05507600.")

    add_heading_custom(doc, "2.2  Model and transfer", level=2)
    add_para(
        doc,
        "The model is an EA-LSTM: static catchment attributes control the input "
        "gate; weather and memory drive the remaining gates. Hidden size is 128, "
        "lookback 90 days, dropout 0.4, forget-gate bias +3.0. Pretraining runs "
        "10 epochs on the 17 donors. Four adaptations are compared on the target: "
        "zero-shot (no local training); Approach A (train only the output head on "
        "the warmup, LSTM frozen); local-from-scratch on the warmup; and "
        "walk-forward evaluation with a head-only refit every 90 days from "
        "2009-01-01, plus online bias correction. Approach B (partial LSTM "
        "unfreeze) was not used in this study.",
        align="justify",
    )
    add_para(
        doc,
        "Walk-forward is the operational test: forecast the next block, score it, "
        "then refit using only discharge that would have been available by that "
        "date. The scored window is 2011–2014 (n = 1,461 days).",
        align="justify",
    )

    add_heading_custom(doc, "2.3  Naive baselines and warnings", level=2)
    add_para(
        doc,
        "Two trivial predictors are scored on the same 2011–2014 window. "
        "Persistence sets Q̂(t) = Q(t−1). Day-of-year climatology is the mean "
        "observed discharge on that calendar day during warmup 2009–2010 only, so "
        "it sees no more local data than the LSTM fine-tune. Skill is NSE, "
        "Kling–Gupta efficiency (KGE), and percent bias (PBIAS).",
        align="justify",
    )
    add_para(
        doc,
        "Warnings are a post-process on the walk-forward hydrograph, not a "
        "separate multi-day weather forecast. Labels at date t with lead L are 1 "
        "if any day in [t+1, t+L] crosses the threshold. Flood thresholds are "
        "Q95 = 3.26 mm d−1 and Q99 = 16.91 mm d−1. Q5 resolves to 0.0 mm d−1 at "
        "this gauge, so those events are low/zero-flow days, not a seasonal "
        "drought. Original probabilities use a Gaussian residual and an "
        "independence product over the lead window. Because consecutive days are "
        "strongly dependent, that product piles probability near 1 as the lead "
        "grows. A second estimator takes the maximum of the 1-day probabilities "
        "in the same window. Discrimination is AUC. Sharpness and calibration "
        "are Brier score and Brier skill score (BSS) against the event base "
        "rate. F1 is reported at 0.5 and at the threshold in {0.05, …, 0.95} "
        "that maximises F1 on the scored window (a diagnostic, not a claimed "
        "operational rule).",
        align="justify",
    )

    add_heading_custom(doc, "3.  Results")
    add_heading_custom(doc, "3.1  Continuous streamflow skill", level=2)
    add_para(
        doc,
        "Table 5 compares the transfer models with the two naive baselines and "
        "the local LSTM. Walk-forward NSE (0.179) is the best continuous score. "
        "Persistence has a high KGE (0.42) because yesterday’s discharge is "
        "correlated with today’s, but its NSE is negative (−0.166): peak errors "
        "on this flashy plains hydrograph more than eat the variance. "
        "Climatology from two years of warmup is worse still (NSE = −2.34). The "
        "local-from-scratch model barely beats a mean forecast (NSE = 0.009) and "
        "is badly wet-biased. Fine-tuning the regional head almost removes "
        "volume bias (PBIAS +4.1%) and lifts KGE from 0.032 (zero-shot) to "
        "0.112. Absolute NSE remains modest next to published continental LSTMs. "
        "What still answers the study question is the ranking: transfer is much "
        "better than local, and transfer beats persistence on NSE.",
        align="justify",
    )
    add_table(
        doc,
        ["Setting", "NSE", "KGE", "PBIAS (%)"],
        [
            ["Next-day persistence",
             _fmt(skill["persistence"]["NSE"]),
             _fmt(skill["persistence"]["KGE"]),
             _pm(skill["persistence"]["PBIAS"])],
            ["Day-of-year climatology (warmup)",
             _fmt(skill["doy_climatology"]["NSE"]),
             _fmt(skill["doy_climatology"]["KGE"]),
             _pm(skill["doy_climatology"]["PBIAS"])],
            ["Local baseline (from-scratch warmup)",
             _fmt(skill["local"]["NSE"]),
             _fmt(skill["local"]["KGE"]),
             _pm(skill["local"]["PBIAS"])],
            ["Zero-shot (pretrained only)",
             _fmt(skill["zero_shot"]["NSE"]),
             _fmt(skill["zero_shot"]["KGE"]),
             _pm(skill["zero_shot"]["PBIAS"])],
            ["Fine-tune, head only (Approach A)",
             _fmt(skill["finetune"]["NSE"]),
             _fmt(skill["finetune"]["KGE"]),
             _pm(skill["finetune"]["PBIAS"])],
            ["Walk-forward (17 × 90-day refits)",
             _fmt(skill["walk_forward"]["NSE"]),
             _fmt(skill["walk_forward"]["KGE"]),
             _pm(skill["walk_forward"]["PBIAS"])],
        ],
    )
    add_caption(
        doc,
        "Table 5. Continuous skill at USGS 05507600, 2011–2014 (n = 1,461 days). "
        "Persistence uses Q(t−1). Climatology is the day-of-year mean of observed "
        "Q during 2009–2010 only.",
    )
    add_figure(
        doc, FIG / "fig_skill_bars.png",
        "Figure 2. NSE and KGE across settings. The dotted line is persistence NSE. "
        "Day-of-year climatology (NSE = −2.34) is annotated rather than plotted on "
        "the same axis.",
    )
    add_figure(
        doc, FIG / "fig_hydrograph_2011_2014.png",
        "Figure 3. Observed discharge versus walk-forward transfer and "
        "local-from-scratch. Transfer follows event timing; the local model does "
        "not. Predictions are clipped at zero for display.",
    )

    add_heading_custom(doc, "3.2  Early-warning scores and calibration", level=2)
    q95_1 = ews["flood_q95_lead1d"]
    q95_3 = ews["flood_q95_lead3d"]
    q95_7 = ews["flood_q95_lead7d"]
    q5_1 = ews["drought_q5_lead1d"]
    add_para(
        doc,
        "Flood Q95 discrimination is strong at every lead (AUC 0.908–0.922). Raw "
        "Brier scores still look as if skill collapses from 0.094 (1-day) to 0.693 "
        "(7-day). BSS against the base rate is negative for the independence "
        "product, and reliability diagrams show why: mean forecast probability at "
        "1-day is 0.29 while the event rate is 0.034, and the 7-day product has "
        "mean probability 0.89. Ranking is intact; calibration is not. Replacing "
        "the product with a window-max of 1-day probabilities cuts 3-day Brier "
        f"from 0.374 to {_fmt(q95_3['window_max']['Brier'])} (BSS "
        f"{_fmt(q95_3['window_max']['BSS'])}) and 7-day Brier from 0.693 to "
        f"{_fmt(q95_7['window_max']['Brier'])} (BSS "
        f"{_fmt(q95_7['window_max']['BSS'])}). F1 at 0.5 is the wrong operating "
        f"point. Sweeping the threshold raises 1-day F1 from 0.286 to "
        f"{_fmt(q95_1['F1_best'])} at probability "
        f"{_fmt(q95_1['threshold'], 2)}.",
        align="justify",
    )
    add_table(
        doc,
        ["Event", "Lead", "AUC", "Brier (prod.)", "BSS (prod.)",
         "Brier (max)", "BSS (max)", "F1@0.5", "F1* (thr)"],
        [
            ["Flood Q95", "1 d",
             _fmt(q95_1["AUC"]), _fmt(q95_1["Brier"]), _fmt(q95_1["BSS"]),
             _fmt(q95_1["window_max"]["Brier"]), _fmt(q95_1["window_max"]["BSS"]),
             _fmt(q95_1["F1@0.5"]),
             f"{_fmt(q95_1['F1_best'])} ({_fmt(q95_1['threshold'], 2)})"],
            ["Flood Q95", "3 d",
             _fmt(q95_3["AUC"]), _fmt(q95_3["Brier"]), _fmt(q95_3["BSS"]),
             _fmt(q95_3["window_max"]["Brier"]), _fmt(q95_3["window_max"]["BSS"]),
             _fmt(q95_3["F1@0.5"]),
             f"{_fmt(q95_3['F1_best'])} ({_fmt(q95_3['threshold'], 2)})"],
            ["Flood Q95", "7 d",
             _fmt(q95_7["AUC"]), _fmt(q95_7["Brier"]), _fmt(q95_7["BSS"]),
             _fmt(q95_7["window_max"]["Brier"]), _fmt(q95_7["window_max"]["BSS"]),
             _fmt(q95_7["F1@0.5"]),
             f"{_fmt(q95_7['F1_best'])} ({_fmt(q95_7['threshold'], 2)})"],
            ["Flood Q99", "1 d",
             _fmt(ews["flood_q99_lead1d"]["AUC"]),
             _fmt(ews["flood_q99_lead1d"]["Brier"]),
             _fmt(ews["flood_q99_lead1d"]["BSS"]),
             _fmt(ews["flood_q99_lead1d"]["window_max"]["Brier"]),
             _fmt(ews["flood_q99_lead1d"]["window_max"]["BSS"]),
             "—", "—"],
            ["Low/zero flow (Q5)", "1 d",
             _fmt(q5_1["AUC"]), _fmt(q5_1["Brier"]), _fmt(q5_1["BSS"]),
             _fmt(q5_1["window_max"]["Brier"]), _fmt(q5_1["window_max"]["BSS"]),
             _fmt(q5_1["F1@0.5"]),
             f"{_fmt(q5_1['F1_best'])} ({_fmt(q5_1['threshold'], 2)})"],
        ],
    )
    add_caption(
        doc,
        "Table 6. Walk-forward warning scores. “Prod.” is the original "
        "independence product; “max” is the window-max of 1-day probabilities. "
        "F1* is the best F1 on a 0.05–0.95 sweep (diagnostic). Q99 F1 is undefined "
        "because too few events exceed any useful probability cutoff. Q5 = 0 mm d−1.",
    )
    add_figure(
        doc, FIG / "fig_reliability_q95.png",
        "Figure 4. Reliability diagrams for flood Q95. Coral: independence product. "
        "Teal: window-max of 1-day probabilities. The product is badly overconfident "
        "at 3- and 7-day leads; window-max pulls the curve back toward the diagonal.",
    )

    add_heading_custom(doc, "3.3  SHAP drivers", level=2)
    add_para(
        doc,
        "GradientExplainer on the fine-tuned model (200 background / 100 sample "
        "sequences) shows precipitation as the most important feature, followed "
        "by vapor pressure, shortwave radiation, day length, Tmin, and Tmax "
        "(Table 7). Static attributes have mean |SHAP| = 0 in this single-basin "
        "explanation because they do not vary within one gauge. That is an "
        "attribution artifact, not evidence that the input gate ignores "
        "catchment properties.",
        align="justify",
    )
    add_table(
        doc,
        ["Feature", "Mean |SHAP|"],
        [
            ["Precipitation (mm/day)", "0.0097"],
            ["Vapor pressure (Pa)", "0.0055"],
            ["Shortwave radiation (W/m²)", "0.0052"],
            ["Day length (s)", "0.0041"],
            ["Tmin (°C)", "0.0036"],
            ["Tmax (°C)", "0.0035"],
            ["Static catchment attributes (all)", "0.0000"],
        ],
    )
    add_caption(doc, "Table 7. Top SHAP drivers for USGS 05507600.")

    add_heading_custom(doc, "3.4  How few donors? A negative result", level=2)
    p17 = prune["17_fulltemp"]
    p8 = prune["8_fulltemp"]
    p8e = prune["8_meanT_ensemble"]
    add_para(
        doc,
        "Table 8 keeps the donor-pruning comparison. Cutting 17 similar donors to "
        "8, with Tmin and Tmax left as separate inputs, lowers walk-forward NSE "
        "from 0.179 to 0.138. Combining that pruning with mean daily temperature "
        "and a three-seed average of daily streamflow predictions recovers some "
        "of the loss (walk-forward NSE 0.165) but still does not beat the "
        "17-donor run. The third row of Table 8 changes three variables at once "
        "(donor count, temperature features, and multi-seed averaging) and "
        "should not be read as a single-factor effect. Forget-gate bias +3.0 is "
        "already the default in every row. Transfer still beats "
        "local-from-scratch after pruning. On this extract, the larger "
        "similar-donor set is the better regional prior.",
        align="justify",
    )
    add_table(
        doc,
        ["Setting", "Donors", "Fine-tune NSE", "Local NSE", "Walk-forward NSE"],
        [
            ["17 donors, Tmin+Tmax, seed 42", "17",
             _fmt(p17["finetune"]["NSE"]), _fmt(p17["local"]["NSE"]),
             _fmt(p17["walk_forward"]["NSE"])],
            ["8 donors, Tmin+Tmax, seed 42", "8",
             _fmt(p8["finetune"]["NSE"]), _fmt(p8["local"]["NSE"]),
             _fmt(p8["walk_forward"]["NSE"])],
            ["8 donors, mean T, 3-seed average", "8",
             _fmt(p8e["finetune"]["NSE"]), _fmt(p8e["local"]["NSE"]),
             _fmt(p8e["walk_forward"]["NSE"])],
        ],
    )
    add_caption(
        doc,
        "Table 8. Sensitivity to donor pruning and related efficiency tweaks "
        "(USGS 05507600, 2011–2014 NSE). The third row changes donor count, "
        "temperature inputs, and seed averaging together.",
    )
    add_figure(
        doc, FIG / "fig_donor_pruning.png",
        "Figure 5. Seventeen similar donors outperform an 8-donor pool on both "
        "fine-tune and walk-forward NSE.",
    )

    add_heading_custom(doc, "4.  Discussion and limits")
    add_para(
        doc,
        "The completed experiment is a one-basin Midwest study on commodity "
        "hardware (Apple M4, about 16 GB RAM). It does not estimate continental "
        "CAMELS skill, snowmelt-regime transfer, or a seven-basin multi-regime "
        "ranking. Those remain available as configs. What it does estimate, with "
        "a full 2011–2014 evaluation, is whether a small similar-donor prior "
        "beats local training and naive baselines when local data are scarce. "
        "It does.",
        align="justify",
    )
    add_para(
        doc,
        "Two caveats belong on a poster. First, warnings are hindcast "
        "post-processing of daily flow, so 3- and 7-day scores are not issued "
        "forecasts. Second, F1* uses the same window it scores; treat it as a "
        "demonstration that 0.5 is a poor cutoff, not as a tuned alarm rule. "
        "Daily 1-day probabilities remain overconfident (BSS = −1.90) even before "
        "compounding. Gaussian residual width is large relative to a 3.4% Q95 "
        "rate. Recalibration of that 1-day probability is future work; it was "
        "not done here because it would require holding out part of 2011–2014 or "
        "retraining.",
        align="justify",
    )

    add_heading_custom(doc, "5.  Conclusion")
    add_para(
        doc,
        "With two years of local discharge and a laptop, seventeen similar Midwest "
        "donors are enough for regional transfer to beat both a local LSTM and "
        "next-day persistence on NSE. Eight donors are worse, not better. Flood "
        "Q95 warnings discriminate well; the ugly raw Brier numbers at long leads "
        "are mostly an independence assumption across consecutive days, not a "
        "collapse of ranking skill. That is a complete, defensible result for a "
        "poster.",
        align="justify",
    )

    add_heading_custom(doc, "Code and reproducibility")
    add_para(
        doc,
        "Midwest run: python scripts/run_midwest_mini.py (configs/midwest_mini/). "
        "Baselines, BSS, threshold sweep, and figures: python "
        "scripts/analysis/poster_from_csvs.py. This Word file: python "
        "scripts/paper/build_short_paper_docx.py. Donor-pruning follow-up: "
        "python scripts/run_midwest_opt.py. Artifacts: results/midwest_mini/ "
        "and docs/paper_images/.",
        align="justify", size=10,
    )

    add_heading_custom(doc, "Selected references")
    refs = [
        "Addor, N., Newman, A. J., Mizukami, N., & Clark, M. P. (2017). The CAMELS data set: catchment attributes and meteorology for large-sample studies. Hydrology and Earth System Sciences, 21, 5293–5313.",
        "Kratzert, F., Klotz, D., Shalev, G., Klambauer, G., Hochreiter, S., & Nearing, G. (2019). Towards learning universal, regional, and local hydrological behaviors via machine learning applied to large-sample datasets. Hydrology and Earth System Sciences, 23, 5089–5110.",
        "Pool, S., Vis, M., & Seibert, J. (2021). Regionalization for ungauged catchments — lessons learned from a comparative large-sample study. Water Resources Research, 57, e2021WR030437.",
        "Gupta, H. V., Kling, H., Yilmaz, K. K., & Martinez, G. F. (2009). Decomposition of the mean squared error and NSE performance criteria. Journal of Hydrology, 377, 80–91.",
        "Brier, G. W. (1950). Verification of forecasts expressed in terms of probability. Monthly Weather Review, 78, 1–3.",
    ]
    for r in refs:
        add_para(doc, r, size=10, space_after=4, first_line=False)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT} ({OUT.stat().st_size / 1024:.0f} KB)")
    return OUT


if __name__ == "__main__":
    build()
